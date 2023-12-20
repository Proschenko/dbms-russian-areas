from openpyxl import Workbook
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse
from django.db import connection
from docx import  Document
from .signals import my_signal
import json
from .models import Area, City, District, Street, ResidentialBuilding, Apartment, Citizen
from .forms import AreaForm, CityForm, DistrictForm, StreetForm, ResidentialBuildingForm, ApartmentForm, CitizenForm, \
    DeleteTypeForm


def index_page(request):
    return render(request, 'indexpage/index.html')


# region report

def export_visible_rows(request,report_type='default'):

    if request.method == 'GET':
        visible_citizen_ids = request.GET.get('visible_citizen_ids')
        report_type = request.GET.get('report_type')

        visible_citizen_ids_list = json.loads(visible_citizen_ids)
        visible_citizen_ids_list = list(map(int, visible_citizen_ids_list))
        citizens = Citizen.objects.filter(id_citizen__in=visible_citizen_ids_list)
        print(visible_citizen_ids_list)
        if report_type == 'word':
            return generate_word_report(citizens)

        elif report_type == 'excel':
            return generate_excel_report(citizens)
        else:
            return HttpResponse("Invalid report type")





def report(request):
    citizens = Citizen.objects.all().order_by('id_citizen')
    context = {
        'citizens': citizens,
    }
    return render(request, 'indexpage/report.html', context)


def generate_word_report(citizens):
    document = Document()
    document.add_heading('Отчет по гражданам', level=1)

    # Создаем таблицу
    table = document.add_table(rows=1, cols=11)  # 11 столбцов, так как у вас 11 полей в модели Citizen
    table.style = 'Table Grid'

    # Добавляем заголовки
    headings = table.rows[0].cells
    headings[0].text = 'ФИО'
    headings[1].text = 'Паспортные данные'
    headings[2].text = 'Номер телефона'
    headings[3].text = 'Дата рождения'
    headings[4].text = 'Пол'
    headings[5].text = 'Квартира'
    headings[6].text = 'Дом'
    headings[7].text = 'Улица'
    headings[8].text = 'Район'
    headings[9].text = 'Город'
    headings[10].text = 'Область'

    # Заполняем таблицу данными
    for citizen in citizens:
        row = table.add_row().cells
        row[0].text = citizen.full_name
        row[1].text = citizen.passport_data
        row[2].text = citizen.phone_number
        row[3].text = str(citizen.date_of_birth)
        row[4].text = 'Мужской' if citizen.gender else 'Женский'
        row[5].text = str(citizen.apartment.apartment_number)
        row[6].text = str(citizen.apartment.residential_building.house_number)
        row[7].text = citizen.apartment.residential_building.street.name_street
        row[8].text = citizen.apartment.residential_building.street.district.name_district
        row[9].text = citizen.apartment.residential_building.street.district.city.name_city
        row[10].text = citizen.apartment.residential_building.street.district.city.area.name_area
        print(row[0].text)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=citizen_report.docx'
    document.save(response)
    return response

def generate_excel_report(citizens):
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.append([
        'ФИО', 'Паспортные данные', 'Номер телефона', 'Дата рождения',
        'Пол', 'Квартира', 'Дом', 'Улица', 'Район', 'Город', 'Область'
    ])

    for citizen in citizens:
        worksheet.append([
            citizen.full_name, citizen.passport_data, citizen.phone_number,
            str(citizen.date_of_birth), 'Мужской' if citizen.gender else 'Женский',
            str(citizen.apartment.apartment_number),
            str(citizen.apartment.residential_building.house_number),
            citizen.apartment.residential_building.street.name_street,
            citizen.apartment.residential_building.street.district.name_district,
            citizen.apartment.residential_building.street.district.city.name_city,
            citizen.apartment.residential_building.street.district.city.area.name_area
        ])
    print("excel good")
    print(worksheet)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=citizen_report.xlsx'
    workbook.save(response)
    return response


# endregion
# region Area
def table_area(request):
    areas = Area.objects.all().order_by('id_area')
    context = {
        'areas': areas,
    }
    return render(request, 'indexpage/table_area.html', context)


def create_area(request):
    if request.method == 'POST':
        form = AreaForm(request.POST)
        if form.is_valid():
            form.save()

            return redirect(table_area)
    else:
        form = AreaForm()
    return render(request, 'indexpage/create_area.html', {'form': form})


def edit_area(request, id_area):
    area = get_object_or_404(Area, pk=id_area)

    if request.method == 'POST':
        form = AreaForm(request.POST, instance=area)
        if form.is_valid():
            form.save()
            my_signal.send(sender=area, request=request)
            return redirect(table_area)
    else:
        form = AreaForm(instance=area)
    return render(request, 'indexpage/edit_area.html', {'form': form, 'area': area})


def delete_area(request, id_area):
    area = get_object_or_404(Area, pk=id_area)
    if request.method == 'POST':
        delete_type_form = DeleteTypeForm(request.POST)
        if delete_type_form.is_valid():
            delete_type = delete_type_form.cleaned_data['delete_type']
            if delete_type == 'CASCADE':
                area.delete()
                return redirect('area_list')

            else:
                # Ваш запрос SQL
                delete_area_sql = f"DELETE FROM indexpage_area WHERE id_area = {id_area}"

                # Выполнение запроса
                with connection.cursor() as cursor:
                    cursor.execute(delete_area_sql)

                return redirect('area_list')

    else:
        delete_type_form = DeleteTypeForm()

    context = {
        'delete_type_form': delete_type_form,
        'area_id': id_area,
        'area_name': area.name_area
    }
    return render(request, 'indexpage/delete_area.html', context)


# endregion

# region City
def table_city(request):
    cities = City.objects.all().order_by('id_city')
    context = {
        'cities': cities,
    }
    return render(request, 'indexpage/table_city.html', context)


def create_city(request):
    if request.method == 'POST':
        form = CityForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect(table_city)
    else:
        form = CityForm()
    return render(request, 'indexpage/create_city.html', {'form': form})


def edit_city(request, id_city):
    city = get_object_or_404(City, pk=id_city)

    if request.method == 'POST':
        form = CityForm(request.POST, instance=city)
        if form.is_valid():
            form.save()
            return redirect(table_city)
    else:
        form = CityForm(instance=city)
    return render(request, 'indexpage/edit_city.html', {'form': form, 'city': city})


def delete_city(request, id_city):
    city = get_object_or_404(City, pk=id_city)
    if request.method == 'POST':
        delete_type_form = DeleteTypeForm(request.POST)
        if delete_type_form.is_valid():
            delete_type = delete_type_form.cleaned_data['delete_type']
            if delete_type == 'CASCADE':
                city.delete()
                return redirect('city_list')
            else:
                # Ваш запрос SQL
                delete_city_sql = f"DELETE FROM indexpage_city WHERE id_city = {id_city}"

                # Выполнение запроса
                with connection.cursor() as cursor:
                    cursor.execute(delete_city_sql)

                return redirect('city_list')
    else:
        delete_type_form = DeleteTypeForm()

    context = {
        'delete_type_form': delete_type_form,
        'city_id': id_city,
        'city_name': city.name_city
    }
    return render(request, 'indexpage/delete_city.html', context)


# endregion

# region District
def table_district(request):
    districts = District.objects.all().order_by('id_district')
    context = {
        'districts': districts,
    }
    return render(request, 'indexpage/table_district.html', context)


def create_district(request):
    if request.method == 'POST':
        form = DistrictForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect(table_district)
    else:
        form = DistrictForm()
    return render(request, 'indexpage/create_district.html', {'form': form})


def edit_district(request, id_district):
    district = get_object_or_404(District, pk=id_district)

    if request.method == 'POST':
        form = DistrictForm(request.POST, instance=district)
        if form.is_valid():
            form.save()
            return redirect(table_district)
    else:
        form = DistrictForm(instance=district)
    return render(request, 'indexpage/edit_district.html', {'form': form, 'district': district})


def delete_district(request, id_district):
    district = get_object_or_404(District, pk=id_district)
    if request.method == 'POST':
        delete_type_form = DeleteTypeForm(request.POST)
        if delete_type_form.is_valid():
            delete_type = delete_type_form.cleaned_data['delete_type']
            if delete_type == 'CASCADE':
                district.delete()
                return redirect('district_list')
            else:
                # Ваш запрос SQL
                delete_district_sql = f"DELETE FROM indexpage_district WHERE id_district = {id_district}"

                # Выполнение запроса
                with connection.cursor() as cursor:
                    cursor.execute(delete_district_sql)

                return redirect('district_list')
    else:
        delete_type_form = DeleteTypeForm()

    context = {
        'delete_type_form': delete_type_form,
        'district_id': id_district,
        'district_name': district.name_district
    }
    return render(request, 'indexpage/delete_district.html', context)


# endregion

# region Street
def table_street(request):
    streets = Street.objects.all().order_by('id_street')
    context = {
        'streets': streets,
    }
    return render(request, 'indexpage/table_street.html', context)


def create_street(request):
    if request.method == 'POST':
        form = StreetForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect(table_street)
    else:
        form = StreetForm()
    return render(request, 'indexpage/create_street.html', {'form': form})


def edit_street(request, id_street):
    street = get_object_or_404(Street, pk=id_street)

    if request.method == 'POST':
        form = StreetForm(request.POST, instance=street)
        if form.is_valid():
            form.save()
            return redirect(table_street)
    else:
        form = StreetForm(instance=street)
    return render(request, 'indexpage/edit_street.html', {'form': form, 'street': street})


def delete_street(request, id_street):
    street = get_object_or_404(Street, pk=id_street)
    if request.method == 'POST':
        delete_type_form = DeleteTypeForm(request.POST)
        if delete_type_form.is_valid():
            delete_type = delete_type_form.cleaned_data['delete_type']
            if delete_type == 'CASCADE':
                street.delete()
                return redirect('street_list')
            else:
                # Ваш запрос SQL
                delete_street_sql = f"DELETE FROM indexpage_street WHERE id_street = {id_street}"

                # Выполнение запроса
                with connection.cursor() as cursor:
                    cursor.execute(delete_street_sql)

                return redirect('street_list')
    else:
        delete_type_form = DeleteTypeForm()

    context = {
        'delete_type_form': delete_type_form,
        'street_id': id_street,
        'street_name': street.name_street
    }
    return render(request, 'indexpage/delete_street.html', context)


# endregion

# region ResidentialBuilding
def table_residentialbuilding(request):
    residential_buildings = ResidentialBuilding.objects.all().order_by('id_residential_building')
    context = {
        'residential_buildings': residential_buildings,
    }
    return render(request, 'indexpage/table_residentialbuilding.html', context)


def create_residentialbuilding(request):
    if request.method == 'POST':
        form = ResidentialBuildingForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect(table_residentialbuilding)
    else:
        form = ResidentialBuildingForm()
    return render(request, 'indexpage/create_residentialbuilding.html', {'form': form})


def edit_residentialbuilding(request, id_residential_building):
    residential_building = get_object_or_404(ResidentialBuilding, pk=id_residential_building)

    if request.method == 'POST':
        form = ResidentialBuildingForm(request.POST, instance=residential_building)
        if form.is_valid():
            form.save()
            return redirect(table_residentialbuilding)
    else:
        form = ResidentialBuildingForm(instance=residential_building)
    return render(request, 'indexpage/edit_residentialbuilding.html',
                  {'form': form, 'residential_building': residential_building})


def delete_residentialbuilding(request, id_residential_building):
    residential_building = get_object_or_404(ResidentialBuilding, pk=id_residential_building)
    if request.method == 'POST':
        delete_type_form = DeleteTypeForm(request.POST)
        if delete_type_form.is_valid():
            delete_type = delete_type_form.cleaned_data['delete_type']
            if delete_type == 'CASCADE':
                residential_building.delete()
                return redirect('residentialbuilding_list')
            else:
                # Ваш запрос SQL
                delete_residential_building_sql = f"DELETE FROM indexpage_residentialbuilding WHERE id_residential_building = {id_residential_building}"

                # Выполнение запроса
                with connection.cursor() as cursor:
                    cursor.execute(delete_residential_building_sql)

                return redirect('residentialbuilding_list')
    else:
        delete_type_form = DeleteTypeForm()

    context = {
        'delete_type_form': delete_type_form,
        'residential_building_id': id_residential_building,
        'residential_building_info': f"{residential_building.house_number}, {residential_building.street.name_street}, {residential_building.street.district.name_district}, {residential_building.street.district.city.name_city}"
    }
    return render(request, 'indexpage/delete_residentialbuilding.html', context)


# endregion

# region Apartment
def table_apartment(request):
    apartments = Apartment.objects.all().order_by('id_apartment')
    context = {
        'apartments': apartments,
    }
    return render(request, 'indexpage/table_apartment.html', context)


def create_apartment(request):
    if request.method == 'POST':
        form = ApartmentForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect(table_apartment)
    else:
        form = ApartmentForm()
    return render(request, 'indexpage/create_apartment.html', {'form': form})


def edit_apartment(request, id_apartment):
    apartment = get_object_or_404(Apartment, pk=id_apartment)

    if request.method == 'POST':
        form = ApartmentForm(request.POST, instance=apartment)
        if form.is_valid():
            form.save()
            return redirect(table_apartment)
    else:
        form = ApartmentForm(instance=apartment)
    return render(request, 'indexpage/edit_apartment.html', {'form': form, 'apartment': apartment})


def delete_apartment(request, id_apartment):
    apartment = get_object_or_404(Apartment, pk=id_apartment)
    if request.method == 'POST':
        delete_type_form = DeleteTypeForm(request.POST)
        if delete_type_form.is_valid():
            delete_type = delete_type_form.cleaned_data['delete_type']
            if delete_type == 'CASCADE':
                apartment.delete()
                return redirect('apartment_list')
            else:
                # Ваш запрос SQL
                delete_apartment_sql = f"DELETE FROM indexpage_apartment WHERE id_apartment = {id_apartment}"

                # Выполнение запроса
                with connection.cursor() as cursor:
                    cursor.execute(delete_apartment_sql)

                return redirect('apartment_list')
    else:
        delete_type_form = DeleteTypeForm()

    context = {
        'delete_type_form': delete_type_form,
        'apartment_id': id_apartment,
        'apartment_info': f"Apartment Number: {apartment.apartment_number}, Building: {apartment.residential_building.house_number}, Street: {apartment.residential_building.street.name_street}, District: {apartment.residential_building.street.district.name_district}, City: {apartment.residential_building.street.district.city.name_city}"
    }
    return render(request, 'indexpage/delete_apartment.html', context)


# endregion

# region Citizen
def table_citizen(request):
    citizens = Citizen.objects.all().order_by('id_citizen')
    context = {
        'citizens': citizens,
    }
    return render(request, 'indexpage/table_citizen.html', context)


def create_citizen(request):
    if request.method == 'POST':
        form = CitizenForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('citizen_list')
    else:
        form = CitizenForm()
    return render(request, 'indexpage/create_citizen.html', {'form': form})


def edit_citizen(request, id_citizen):
    citizen = get_object_or_404(Citizen, pk=id_citizen)

    if request.method == 'POST':
        form = CitizenForm(request.POST, instance=citizen)
        if form.is_valid():
            form.save()
            return redirect('citizen_list')
    else:
        form = CitizenForm(instance=citizen)
    return render(request, 'indexpage/edit_citizen.html', {'form': form, 'citizen': citizen})


def delete_citizen(request, id_citizen):
    citizen = get_object_or_404(Citizen, pk=id_citizen)
    if request.method == 'POST':
        delete_type_form = DeleteTypeForm(request.POST)
        if delete_type_form.is_valid():
            delete_type = delete_type_form.cleaned_data['delete_type']
            if delete_type == 'CASCADE':
                citizen.delete()
                return redirect('citizen_list')
            else:
                # Ваш запрос SQL
                delete_citizen_sql = f"DELETE FROM indexpage_citizen WHERE id_citizen = {id_citizen}"

                # Выполнение запроса
                with connection.cursor() as cursor:
                    cursor.execute(delete_citizen_sql)

                return redirect('citizen_list')
    else:
        delete_type_form = DeleteTypeForm()

    context = {
        'delete_type_form': delete_type_form,
        'citizen_id': id_citizen,
        'citizen_info': f"Full Name: {citizen.full_name}, Passport Data: {citizen.passport_data}, Phone Number: {citizen.phone_number}, Date of Birth: {citizen.date_of_birth}, Gender: {'Male' if citizen.gender else 'Female'}, Apartment: {citizen.apartment}"
    }
    return render(request, 'indexpage/delete_citizen.html', context)
# endregion
